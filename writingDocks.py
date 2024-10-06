import json
import os
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
keywords = ["candle", "red", "white", "doji", "yellow", "green", "blue", "black", "downward", "up", "down", "flat", "upward", "spinning top", "cross", "gapped", "MA", "SRSI", "MACD", "DM", "TBB", "BBB", "auto-wave", "declining", "rising", "flattening", "sideways", "resistance", "support", "chart", "permission", "downside", "sideways up", "sideways down", "squeeze", "Bollinger Band", "noise", "gap down", "bounce", "test", "shifted", "crossing", "level", "extreme", "cocked", "trigger line","nail", "21MA", "233MA", "377MA", "55MA", "89MA", "144MA", "volatility", "sideways movement", "Bollinger Band Squeeze", "sideways to higher", "opened", "close", "bounced", "bands", "magnetic", "moving", "higher", "lower", "Quarterly", "purple line", "double top", "capital M", "34MA", "20-level", "Reverse Head & Shoulders", "green line", "down auto-wave", "1-3 years", "moving sideways"]


def highlight_keywords(text, run):
    words = text.split()
    current_run = None
    for word in words:
        if any(keyword.lower() in word.lower() for keyword in keywords):
            current_run = run.add_run(word + ' ')
            current_run.bold = True
            current_run.font.color.rgb = RGBColor(0, 0, 0)  # Red color
        else:
            current_run = run.add_run(word + ' ')
        current_run.font.name = 'Arial'
        current_run.font.size = Pt(12)

def createDoc(ofWhat, withWhat, month, rootDir):
    doc = Document()
    doc.add_heading(ofWhat, level=1)

    for item in withWhat:
        for key, value in item.items():
            p = doc.add_paragraph()
            run = p.add_run(f'{month} {key}: ')
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(12)

            highlight_keywords(value, p)

    doc.save(f'{rootDir}/{ofWhat}.docx')

def createDirIfNotExists(directoryPath):
    if not os.path.exists(directoryPath):
        os.makedirs(directoryPath)

def sortDictsByKeys(saveLocation, data):
    # createDirIfNotExists(saveLocation)
    for month, categories in data.items():
        folderName = f"{month}_ChartWatch"
        createDirIfNotExists(saveLocation + '/' + folderName)
        for symbol, schedules in categories.items():
            docxDir = saveLocation + '/' + folderName + '/' + symbol
            createDirIfNotExists(docxDir)
            for scheduleType, entries in schedules.items():
                sortedEntries = sorted(entries, key=lambda x: int(next(iter(x))))
                schedules[scheduleType] = sortedEntries
                createDoc(scheduleType, sortedEntries, month, docxDir)

# Start processing

def callFromMaster(saveLocation):
    with open('stockData.json', 'r') as file:
        stockData = json.load(file)
    sortDictsByKeys(saveLocation, stockData)
    os.remove('stockData.json')
